import io
from datetime import date

from fastapi.testclient import TestClient

from app.main import create_app
from app.models import Application, Base, User
from app.security import hash_password


class MemoryObjects:
    def __init__(self) -> None:
        self.objects: dict[str, bytes] = {}

    def put(self, key: str, stream, length: int) -> None:
        chunks = []
        remaining = length
        while remaining != 0:
            chunk = stream.read(65536 if remaining < 0 else min(65536, remaining))
            if not chunk:
                break
            chunks.append(chunk)
            if remaining > 0:
                remaining -= len(chunk)
        self.objects[key] = b"".join(chunks)

    def open(self, key: str):
        return io.BytesIO(self.objects[key])

    def delete(self, key: str) -> None:
        self.objects.pop(key, None)


def setup() -> tuple[TestClient, str]:
    app = create_app(
        "sqlite+pysqlite:///:memory:", cookie_secure=True, object_store=MemoryObjects()
    )
    Base.metadata.create_all(app.state.database.engine)
    with app.state.database.session() as db:
        owner = User(
            username="owner",
            password_hash=hash_password("approval officer password"),
            role="approval_officer",
        )
        other = User(
            username="other",
            password_hash=hash_password("approval officer password"),
            role="approval_officer",
        )
        db.add_all([owner, other])
        db.flush()
        application = Application(
            borrower_type="corporate",
            borrower_name="示例企业",
            product="经营贷",
            application_date=date(2026, 8, 7),
            owner_id=owner.id,
        )
        db.add(application)
        db.commit()
        application_id = application.id
    return TestClient(app, base_url="https://testserver"), application_id


def login(client: TestClient, username: str) -> dict[str, str]:
    response = client.post(
        "/api/v1/auth/login",
        headers={"Origin": "https://testserver"},
        json={"username": username, "password": "approval officer password"},
    )
    assert response.status_code == 204
    return {"X-CSRF-Token": client.cookies["icrm_csrf"]}


def upload(client: TestClient, application_id: str, headers: dict[str, str], key: str = "one"):
    return client.post(
        f"/api/v1/applications/{application_id}/documents",
        headers={**headers, "Idempotency-Key": key},
        files={"file": ("sample.pdf", b"%PDF-1.7\nsynthetic", "application/pdf")},
    )


def test_upload_returns_document_and_waiting_job_and_deduplicates() -> None:
    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        created = upload(client, application_id, csrf)
        assert created.status_code == 202
        assert created.json()["document"]["filename"] == "sample.pdf"
        assert created.json()["job"]["status"] == "waiting"
        assert [step["name"] for step in created.json()["job"]["steps"]] == [
            "validation",
            "parsing_ocr",
            "structure_extraction",
            "seal_detection",
            "classification",
            "candidate_extraction",
        ]

        replay = upload(client, application_id, csrf)
        assert replay.status_code == 200
        assert replay.json()["document"]["id"] == created.json()["document"]["id"]
        mismatch = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "one"},
            files={"file": ("other.pdf", b"%PDF-1.7\nother", "application/pdf")},
        )
        assert mismatch.status_code == 409

        duplicate = upload(client, application_id, csrf, "two")
        assert duplicate.status_code == 200
        assert duplicate.json()["document"]["id"] == created.json()["document"]["id"]
        reused_duplicate_key = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "two"},
            files={"file": ("changed.pdf", b"%PDF-1.7\nchanged", "application/pdf")},
        )
        assert reused_duplicate_key.status_code == 409
        same_bytes_changed_metadata = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "one"},
            files={"file": ("renamed.pdf", b"%PDF-1.7\nsynthetic", "text/plain")},
        )
        assert same_bytes_changed_metadata.status_code == 409
        assert [step["status"] for step in duplicate.json()["job"]["steps"]] == [
            "waiting",
            "waiting",
            "not_applicable",
            "waiting",
            "not_applicable",
            "waiting",
        ]


def test_upload_enforces_limits_and_owner_boundary() -> None:
    client, application_id = setup()
    client.app.state.document_limits.max_material_bytes = 8
    with client:
        csrf = login(client, "owner")
        oversized = upload(client, application_id, csrf)
        assert oversized.status_code == 202
        assert oversized.json()["job"]["status"] == "failed"
        assert oversized.json()["job"]["error_code"] == "material_size_limit_exceeded"
        assert oversized.json()["document"]["id"]
        assert len(oversized.json()["job"]["steps"]) == 6
        replay = upload(client, application_id, csrf)
        assert replay.status_code == 200
        assert replay.json()["document"]["id"] == oversized.json()["document"]["id"]
        client.app.state.document_limits.max_application_materials = 1
        replay_at_capacity = upload(client, application_id, csrf)
        assert replay_at_capacity.status_code == 200
        duplicate_at_capacity = upload(client, application_id, csrf, "duplicate-at-capacity")
        assert duplicate_at_capacity.status_code == 200
        assert duplicate_at_capacity.json()["document"]["id"] == oversized.json()["document"]["id"]
        another = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "another-oversized"},
            files={"file": ("other.pdf", b"%PDF-1.7\ndifferent", "application/pdf")},
        )
        assert another.status_code == 413

    client.cookies.clear()
    with client:
        csrf = login(client, "other")
        assert upload(client, application_id, csrf).status_code == 404
        assert client.get(f"/api/v1/applications/{application_id}/documents").status_code == 404


def test_oversized_outcome_respects_application_size_limit() -> None:
    client, application_id = setup()
    client.app.state.document_limits.max_material_bytes = 8
    client.app.state.document_limits.max_application_bytes = 25
    with client:
        csrf = login(client, "owner")
        assert upload(client, application_id, csrf).status_code == 202
        another = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "another-oversized-size"},
            files={"file": ("other.pdf", b"%PDF-1.7\ndifferent", "application/pdf")},
        )
        assert another.status_code == 413


def test_manual_handling_does_not_hide_another_waiting_material() -> None:
    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        assert upload(client, application_id, csrf).status_code == 202
        manual = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "legacy-with-active"},
            files={"file": ("legacy.doc", b"legacy", "application/msword")},
        )
        assert manual.status_code == 202
        application = client.get(f"/api/v1/applications/{application_id}").json()
        assert application["lifecycle_state"] == "processing"


def test_unsupported_extension_is_explicit_manual_handling() -> None:
    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        response = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "legacy"},
            files={"file": ("legacy.doc", b"legacy", "application/msword")},
        )
        assert response.status_code == 202
        assert response.json()["job"]["status"] == "manual_handling"
        assert response.json()["job"]["error_code"] == "unsupported_legacy_office"
        application = client.get(f"/api/v1/applications/{application_id}").json()
        assert application["lifecycle_state"] == "pending_review"


def test_retry_requires_reason_and_owner() -> None:
    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        response = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "archive"},
            files={"file": ("archive.zip", b"PK\x03\x04", "application/zip")},
        )
        job_id = response.json()["job"]["id"]
        assert (
            client.post(
                f"/api/v1/jobs/{job_id}/retry", headers=csrf, json={"reason": ""}
            ).status_code
            == 422
        )
        retry_headers = {**csrf, "Idempotency-Key": "retry-archive"}
        retry_payload = {"reason": "已重新确认材料格式", "selected_steps": ["validation"]}
        retried = client.post(
            f"/api/v1/jobs/{job_id}/retry",
            headers=retry_headers,
            json=retry_payload,
        )
        assert retried.status_code == 200
        assert retried.json()["retry_reason"] == "已重新确认材料格式"
        assert retried.json()["status"] == "waiting"
        application = client.get(f"/api/v1/applications/{application_id}").json()
        assert application["lifecycle_state"] == "processing"
        replay = client.post(
            f"/api/v1/jobs/{job_id}/retry",
            headers=retry_headers,
            json=retry_payload,
        )
        assert replay.status_code == 200
        assert replay.json() == retried.json()
        mismatch = client.post(
            f"/api/v1/jobs/{job_id}/retry",
            headers=retry_headers,
            json={"reason": "不同理由", "selected_steps": ["validation"]},
        )
        assert mismatch.status_code == 409
        rejected = client.post(
            f"/api/v1/jobs/{job_id}/retry",
            headers={**csrf, "Idempotency-Key": "invalid-step"},
            json={"reason": "尝试运行后续步骤", "selected_steps": ["parsing_ocr"]},
        )
        assert rejected.status_code == 409

    client.cookies.clear()
    with client:
        other_csrf = login(client, "other")
        document_id = response.json()["document"]["id"]
        assert client.get(f"/api/v1/documents/{document_id}/jobs").status_code == 404
        unauthorized_retry = client.post(
            f"/api/v1/jobs/{job_id}/retry",
            headers={**other_csrf, "Idempotency-Key": "other-retry"},
            json=retry_payload,
        )
        assert unauthorized_retry.status_code == 404


def test_owner_reads_versioned_output_and_reruns_selected_parser_steps() -> None:
    import pymupdf

    from app.parsing import Analysis
    from app.worker import process_one

    pdf = pymupdf.open()
    page = pdf.new_page(width=300, height=200)
    page.insert_text((30, 50), "Version one remains reviewable")
    content = pdf.tobytes()

    class Engine:
        version = "test-model-1"

        def analyze(self, image: bytes, *, run_ocr: bool) -> Analysis:
            assert not run_ocr
            return Analysis()

    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        uploaded = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "parsed-pdf"},
            files={"file": ("parsed.pdf", content, "application/pdf")},
        )
        assert process_one(
            client.app.state.database, client.app.state.object_store, "parser", Engine()
        )
        document_id = uploaded.json()["document"]["id"]
        outputs = client.get(f"/api/v1/documents/{document_id}/outputs")
        assert outputs.status_code == 200
        assert outputs.json()[0]["version"] == 1
        assert outputs.json()[0]["pages"][0]["blocks"][0]["text"] == (
            "Version one remains reviewable"
        )
        preview = client.get(f"/api/v1/documents/{document_id}/pages/1/image")
        assert preview.status_code == 200
        assert preview.headers["content-type"] == "image/png"
        from PIL import Image

        assert Image.open(io.BytesIO(preview.content)).size == (600, 400)

        job_id = uploaded.json()["job"]["id"]
        rerun = client.post(
            f"/api/v1/jobs/{job_id}/retry",
            headers={**csrf, "Idempotency-Key": "parser-rerun"},
            json={
                "reason": "使用固定的新模型重新解析",
                "selected_steps": ["parsing_ocr", "seal_detection"],
            },
        )
        assert rerun.status_code == 200
        assert process_one(
            client.app.state.database, client.app.state.object_store, "parser", Engine()
        )
        versions = client.get(f"/api/v1/documents/{document_id}/outputs").json()
        assert [output["version"] for output in versions] == [1, 2]
        assert versions[0]["pages"][0]["blocks"][0]["text"] == (
            "Version one remains reviewable"
        )


def test_owner_confirms_seal_candidate_and_records_signature_presence_manually() -> None:
    from PIL import Image

    from app.parsing import Analysis, BlockResult, SealResult
    from app.worker import process_one

    image = io.BytesIO()
    Image.new("RGB", (120, 80), "white").save(image, format="PNG")

    class Engine:
        version = "seal-model-1"

        def analyze(self, content: bytes, *, run_ocr: bool) -> Analysis:
            assert run_ocr
            return Analysis(
                blocks=(BlockResult(0, "paragraph", "公开样例", (10, 10, 60, 30), "ocr", 0.9),),
                seals=(SealResult("印章文字候选", (60, 30, 110, 75), 0.8),),
            )

    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        uploaded = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "seal-image"},
            files={"file": ("seal.png", image.getvalue(), "image/png")},
        )
        assert process_one(
            client.app.state.database, client.app.state.object_store, "parser", Engine()
        )
        document_id = uploaded.json()["document"]["id"]
        output = client.get(f"/api/v1/documents/{document_id}/outputs").json()[0]
        seal = output["pages"][0]["seals"][0]
        seal_review = client.post(
            f"/api/v1/document-outputs/{output['id']}/reviews",
            headers={**csrf, "Idempotency-Key": "confirm-seal"},
            json={
                "kind": "seal_presence",
                "status": "present",
                "seal_candidate_id": seal["id"],
                "reason": "审批人员查看原页后确认存在印章区域",
            },
        )
        assert seal_review.status_code == 201
        assert seal_review.json()["kind"] == "seal_presence"
        signature_review = client.post(
            f"/api/v1/document-outputs/{output['id']}/reviews",
            headers={**csrf, "Idempotency-Key": "confirm-signature"},
            json={
                "kind": "signature_presence",
                "status": "present",
                "seal_candidate_id": None,
                "reason": "审批人员人工查看原页后确认存在签字",
            },
        )
        assert signature_review.status_code == 201
        reviews = client.get(f"/api/v1/document-outputs/{output['id']}/reviews").json()
        assert [review["kind"] for review in reviews] == [
            "seal_presence",
            "signature_presence",
        ]
        assert all("authentic" not in str(review).lower() for review in reviews)


def test_output_preview_and_review_resources_respect_ownership_and_ordering() -> None:
    import pymupdf

    from app.parsing import Analysis
    from app.worker import process_one

    pdf = pymupdf.open()
    page = pdf.new_page(width=300, height=200)
    page.insert_text((30, 50), "Owned synthetic statement")
    content = pdf.tobytes()

    class Engine:
        version = "test-model-1"

        def analyze(self, image: bytes, *, run_ocr: bool) -> Analysis:
            assert not run_ocr
            return Analysis()

    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        uploaded = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "ownership-pdf"},
            files={"file": ("owned.pdf", content, "application/pdf")},
        )
        assert process_one(
            client.app.state.database, client.app.state.object_store, "parser", Engine()
        )
        document_id = uploaded.json()["document"]["id"]
        output = client.get(f"/api/v1/documents/{document_id}/outputs").json()[0]

    client.cookies.clear()
    with client:
        other_csrf = login(client, "other")
        assert (
            client.get(f"/api/v1/documents/{document_id}/outputs").status_code == 404
        )
        assert (
            client.get(f"/api/v1/documents/{document_id}/pages/1/image").status_code == 404
        )
        assert (
            client.get(f"/api/v1/document-outputs/{output['id']}/reviews").status_code == 404
        )
        assert (
            client.post(
                f"/api/v1/document-outputs/{output['id']}/reviews",
                headers={**other_csrf, "Idempotency-Key": "other-review"},
                json={
                    "kind": "signature_presence",
                    "status": "present",
                    "seal_candidate_id": None,
                    "reason": "无权访问",
                },
            ).status_code
            == 404
        )


def test_review_replay_is_idempotent_and_ordered_by_creation() -> None:
    from PIL import Image

    from app.parsing import Analysis, BlockResult, SealResult
    from app.worker import process_one

    image = io.BytesIO()
    Image.new("RGB", (120, 80), "white").save(image, format="PNG")

    class Engine:
        version = "seal-model-1"

        def analyze(self, content: bytes, *, run_ocr: bool) -> Analysis:
            assert run_ocr
            return Analysis(
                blocks=(BlockResult(0, "paragraph", "公开样例", (10, 10, 60, 30), "ocr", 0.9),),
                seals=(SealResult("印章文字候选", (60, 30, 110, 75), 0.8),),
            )

    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        uploaded = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "review-idem-image"},
            files={"file": ("seal.png", image.getvalue(), "image/png")},
        )
        assert process_one(
            client.app.state.database, client.app.state.object_store, "parser", Engine()
        )
        document_id = uploaded.json()["document"]["id"]
        output = client.get(f"/api/v1/documents/{document_id}/outputs").json()[0]
        payload = {
            "kind": "signature_presence",
            "status": "present",
            "seal_candidate_id": None,
            "reason": "人工查看原页后确认签字存在",
        }
        review_headers = {**csrf, "Idempotency-Key": "sig-review"}
        created = client.post(
            f"/api/v1/document-outputs/{output['id']}/reviews",
            headers=review_headers,
            json=payload,
        )
        assert created.status_code == 201
        replay = client.post(
            f"/api/v1/document-outputs/{output['id']}/reviews",
            headers=review_headers,
            json=payload,
        )
        assert replay.status_code == 201
        assert replay.json()["id"] == created.json()["id"]
        mismatch = client.post(
            f"/api/v1/document-outputs/{output['id']}/reviews",
            headers=review_headers,
            json={**payload, "status": "absent"},
        )
        assert mismatch.status_code == 409
        first = client.post(
            f"/api/v1/document-outputs/{output['id']}/reviews",
            headers={**csrf, "Idempotency-Key": "sig-review-first"},
            json=payload,
        )
        assert first.status_code == 201
        reviews = client.get(f"/api/v1/document-outputs/{output['id']}/reviews").json()
        assert [review["id"] for review in reviews] == [
            created.json()["id"],
            first.json()["id"],
        ]


def test_retrying_unchanged_unsupported_material_keeps_manual_outcome() -> None:
    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        uploaded = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "legacy-retry"},
            files={"file": ("legacy.doc", b"legacy", "application/msword")},
        )
        job_id = uploaded.json()["job"]["id"]
        retried = client.post(
            f"/api/v1/jobs/{job_id}/retry",
            headers={**csrf, "Idempotency-Key": "retry-legacy"},
            json={"reason": "人工要求重试", "selected_steps": ["validation"]},
        )
        assert retried.status_code == 200
        from app.worker import process_one

        assert process_one(client.app.state.database, client.app.state.object_store, "worker")
        status = client.get(f"/api/v1/documents/{uploaded.json()['document']['id']}/jobs").json()
        assert status[-1]["status"] == "manual_handling"
        assert status[-1]["error_code"] == "unsupported_legacy_office"


def test_docx_upload_runs_parsing_without_seal_detection() -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("经营情况说明")
    buffer = io.BytesIO()
    document.save(buffer)

    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        created = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "docx"},
            files={
                "file": (
                    "statement.docx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert created.status_code == 202
        steps = {
            step["name"]: step["status"] for step in created.json()["job"]["steps"]
        }
        assert steps["validation"] == "waiting"
        assert steps["parsing_ocr"] == "waiting"
        assert steps["seal_detection"] == "not_applicable"
        assert steps["structure_extraction"] == "not_applicable"


def test_owner_downloads_original_material_and_boundary_is_enforced() -> None:
    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        created = upload(client, application_id, csrf)
        document_id = created.json()["document"]["id"]
        download = client.get(f"/api/v1/documents/{document_id}/download")
        assert download.status_code == 200
        assert download.content == b"%PDF-1.7\nsynthetic"
        assert download.headers["content-type"] == "application/pdf"
        assert "attachment" in download.headers["content-disposition"]

        with client.app.state.database.session() as db:
            from app.models import Document as DocumentModel

            object_key = db.get(DocumentModel, document_id).object_key
        client.app.state.object_store.delete(object_key)
        assert (
            client.get(f"/api/v1/documents/{document_id}/download").status_code == 404
        )

    client.cookies.clear()
    with client:
        csrf = login(client, "other")
        assert (
            client.get(f"/api/v1/documents/{document_id}/download").status_code == 404
        )


def test_structured_outputs_expose_format_and_native_locators_without_pages() -> None:
    import openpyxl

    from app.parsed_outputs import store_parsed_output
    from app.structured import parse_structured

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "流水明细"
    sheet.append(["日期", "金额"])
    sheet.append(["2026-08-01", 1234.5])
    buffer = io.BytesIO()
    workbook.save(buffer)

    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        uploaded = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "xlsx-output"},
            files={
                "file": (
                    "statement.xlsx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )
        document_id = uploaded.json()["document"]["id"]
        parsed = parse_structured(
            "statement.xlsx", io.BytesIO(buffer.getvalue())
        )
        with client.app.state.database.session() as db:
            store_parsed_output(db, document_id, parsed)
            db.commit()

        outputs = client.get(f"/api/v1/documents/{document_id}/outputs")
        assert outputs.status_code == 200
        output = outputs.json()[0]
        assert output["format"] == "xlsx"
        assert output["status"] == "success"
        page = output["pages"][0]
        assert page["number"] is None
        assert page["width"] is None
        block = page["blocks"][0]
        assert block["kind"] == "table"
        assert block["locator"] == {
            "kind": "xlsx",
            "paragraph_path": None,
            "sheet": "流水明细",
            "cell_range": "A1:B2",
            "cell": None,
            "row": None,
            "column": None,
            "column_name": None,
            "encoding": None,
            "heading_path": None,
            "line_start": None,
            "line_end": None,
        }
        assert block["cells"][0]["locator"]["cell"] == "A1"
        assert block["cells"][3]["text"] == "1234.5"
        # Structured formats never get a page image endpoint.
        assert (
            client.get(f"/api/v1/documents/{document_id}/pages/1/image").status_code
            == 404
        )


def test_structured_rerun_selects_only_the_parsing_step() -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("经营情况说明")
    buffer = io.BytesIO()
    document.save(buffer)

    client, application_id = setup()
    with client:
        csrf = login(client, "owner")
        uploaded = client.post(
            f"/api/v1/applications/{application_id}/documents",
            headers={**csrf, "Idempotency-Key": "docx-rerun"},
            files={
                "file": (
                    "statement.docx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        job_id = uploaded.json()["job"]["id"]
        with client.app.state.database.session() as db:
            from app.document_jobs import claim_next_job, finish_job

            claimed = claim_next_job(db, "api-test")
            finish_job(db, claimed, "success", claim_token=claimed.claim_token)
        rerun = client.post(
            f"/api/v1/jobs/{job_id}/retry",
            headers={**csrf, "Idempotency-Key": "docx-rerun-1"},
            json={"reason": "重新解析", "selected_steps": ["parsing_ocr"]},
        )
        assert rerun.status_code == 200
        rejected = client.post(
            f"/api/v1/jobs/{job_id}/retry",
            headers={**csrf, "Idempotency-Key": "docx-rerun-2"},
            json={"reason": "印章步骤不适用", "selected_steps": ["seal_detection"]},
        )
        assert rejected.status_code == 409
