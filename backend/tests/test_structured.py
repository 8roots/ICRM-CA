import io

import pytest

from app.parsing import Locator, parse_material
from app.structured import (
    decode_text,
    parse_csv,
    parse_structured,
)


def docx_bytes() -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("企业概况", level=1)
    document.add_paragraph("示例企业成立于2015年，主营零售。")
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "科目"
    table.cell(0, 1).text = "金额"
    table.cell(0, 2).text = "期间"
    table.cell(1, 0).text = "营业收入"
    table.cell(1, 1).text = "1,234,567.89"
    table.cell(1, 2).text = "2025"
    table.cell(2, 0).text = "净利润"
    table.cell(2, 1).text = "98,765.43"
    table.cell(2, 2).text = "2025"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def xlsx_bytes() -> bytes:
    import openpyxl

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "流水明细"
    sheet.append(["日期", "交易对手", "收支", "金额"])
    sheet.append(["2026-08-01", "甲公司", "支出", 1234.5])
    sheet.append(["2026-08-02", "乙公司", "收入", 5000])
    sheet.append(["2026-08-03", "丙公司", "收入", 88.25])
    sheet.merge_cells("A5:C5")
    sheet["A5"] = "月度合计见下页"
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def markdown_text() -> str:
    return """# 一、企业概况

示例企业成立于2015年。

## 银行流水

| 日期 | 金额 |
| --- | ---: |
| 2026-08-01 | 1,234.56 |
| 2026-08-02 | -500.00 |
"""


def test_docx_returns_paragraph_and_table_blocks_with_body_paths() -> None:
    parsed = parse_structured("statement.docx", io.BytesIO(docx_bytes()))

    assert parsed.parser_version.startswith("python-docx-")
    assert parsed.model_version == "none"
    assert parsed.status == "success"
    page = parsed.pages[0]
    assert (page.number, page.width, page.height) == (None, None, None)
    kinds = [block.kind for block in page.blocks]
    assert kinds == ["heading", "paragraph", "table"]
    assert [block.locator.paragraph_path for block in page.blocks] == [
        "body/1",
        "body/2",
        "body/3",
    ]
    assert page.blocks[1].text == "示例企业成立于2015年，主营零售。"
    assert page.blocks[0].locator == Locator(kind="docx", paragraph_path="body/1")
    table = page.blocks[2]
    assert table.extraction_method == "docx_text"
    assert [(cell.row, cell.column, cell.text) for cell in table.cells[:3]] == [
        (1, 1, "科目"),
        (1, 2, "金额"),
        (1, 3, "期间"),
    ]
    assert (table.cells[6].row, table.cells[6].text) == (3, "净利润")


def test_xlsx_preserves_rows_columns_merged_cells_and_amounts() -> None:
    parsed = parse_structured("statement.xlsx", io.BytesIO(xlsx_bytes()))

    assert parsed.parser_version.startswith("openpyxl-")
    page = parsed.pages[0]
    assert (page.number, page.width, page.height) == (None, None, None)
    assert len(page.blocks) == 1
    table = page.blocks[0]
    assert table.kind == "table"
    assert table.extraction_method == "xlsx_text"
    assert table.locator == Locator(kind="xlsx", sheet="流水明细", cell_range="A1:D5")
    cells = {(cell.row, cell.column): cell for cell in table.cells}
    assert len(cells) == 4 * 5
    assert cells[(2, 1)].text == "2026-08-01"
    assert cells[(2, 2)].text == "甲公司"
    assert cells[(2, 4)].text == "1234.5"
    assert cells[(3, 4)].text == "5000"
    assert cells[(4, 4)].text == "88.25"
    assert cells[(2, 1)].locator == Locator(kind="xlsx", sheet="流水明细", cell="A2")
    assert cells[(4, 4)].locator == Locator(kind="xlsx", sheet="流水明细", cell="D4")
    # Merged A5:C5 expands into every covered cell; D5 stays empty.
    for column in (1, 2, 3):
        assert cells[(5, column)].text == "月度合计见下页"
    assert cells[(5, 4)].text == ""


def test_xlsx_skips_empty_sheets() -> None:
    import openpyxl

    workbook = openpyxl.Workbook()
    workbook.active.title = "空表"
    buffer = io.BytesIO()
    workbook.save(buffer)

    parsed = parse_structured("empty.xlsx", io.BytesIO(buffer.getvalue()))

    assert parsed.status == "success"
    assert parsed.pages[0].blocks == ()


@pytest.mark.parametrize(
    ("content", "encoding", "expected"),
    [
        (b"column,value\n1,2", "utf-8", [("1", "column"), ("2", "value")]),
        (b"\xef\xbb\xbfcolumn,value\n1,2", "utf-8", [("1", "column"), ("2", "value")]),
        (
            "列,值\n甲,100\n乙,200".encode("gb18030"),
            "gb18030",
            [("甲", "列"), ("乙", "列")],
        ),
    ],
)
def test_csv_encoding_handling_and_row_column_references(
    content: bytes, encoding: str, expected: list[tuple[str, str]]
) -> None:
    parsed = parse_structured("transactions.csv", io.BytesIO(content))

    assert parsed.parser_version == "icrm-csv-1"
    page = parsed.pages[0]
    assert (page.number, page.width, page.height) == (None, None, None)
    assert len(page.blocks) == 1
    table = page.blocks[0]
    assert table.locator.kind == "csv"
    assert table.locator.encoding == encoding
    cells = {(cell.row, cell.column): cell for cell in table.cells}
    for row_index in range(1, 3):
        for column_index in range(1, 3):
            cell = cells[(row_index, column_index)]
            assert cell.locator.row == row_index
            assert cell.locator.column == column_index


def test_csv_header_columns_are_named_and_rows_numbered() -> None:
    blocks = parse_csv(b"date,amount\n2026-08-01,1234.5\n2026-08-02,-50\n")

    cells = {(cell.row, cell.column): cell for cell in blocks[0].cells}
    assert cells[(1, 1)].locator.column_name == "date"
    assert cells[(1, 2)].locator.column_name == "amount"
    assert cells[(2, 1)].text == "2026-08-01"
    assert cells[(2, 2)].text == "1234.5"
    assert cells[(3, 2)].text == "-50"
    assert cells[(3, 1)].locator.column_name == "date"


def test_markdown_heading_path_line_ranges_and_pipe_table() -> None:
    parsed = parse_structured("notes.md", io.BytesIO(markdown_text().encode()))

    assert parsed.parser_version == "icrm-markdown-1"
    blocks = parsed.pages[0].blocks
    assert [(block.kind, block.text) for block in blocks] == [
        ("heading", "一、企业概况"),
        ("paragraph", "示例企业成立于2015年。"),
        ("heading", "银行流水"),
        ("table", "日期"),
    ]
    assert blocks[0].locator == Locator(
        kind="markdown", heading_path="一、企业概况", line_start=1, line_end=1
    )
    assert blocks[1].locator.line_start == 3
    assert blocks[2].locator.heading_path == "一、企业概况 / 银行流水"
    table = blocks[3]
    assert table.locator.line_start == 7
    assert table.locator.line_end == 10
    cells = {(cell.row, cell.column): cell for cell in table.cells}
    assert cells[(1, 1)].text == "日期"
    assert cells[(2, 2)].text == "1,234.56"
    assert cells[(3, 2)].text == "-500.00"
    assert cells[(3, 1)].locator is None


def test_structured_outputs_never_invent_page_numbers() -> None:
    for filename, content in [
        ("a.docx", docx_bytes()),
        ("b.xlsx", xlsx_bytes()),
        ("c.csv", b"a,b\n1,2"),
        ("d.md", markdown_text().encode()),
    ]:
        parsed = parse_structured(filename, io.BytesIO(content))
        for page in parsed.pages:
            assert page.number is None
            assert page.width is None
            assert page.height is None


def test_page_parser_rejects_structured_extensions() -> None:
    class Engine:
        version = "fake"

        def analyze(self, image: bytes, *, run_ocr: bool):
            raise AssertionError("structured formats must not reach the image engine")

    with pytest.raises(ValueError):
        parse_material("a.docx", io.BytesIO(docx_bytes()), Engine())


def test_structured_parser_rejects_page_extensions() -> None:
    with pytest.raises(ValueError):
        parse_structured("a.pdf", io.BytesIO(b"%PDF-1.7"))


def test_decode_text_rejects_undecodable_content() -> None:
    with pytest.raises(UnicodeDecodeError):
        decode_text(b"\xff\xfe\xfd")
