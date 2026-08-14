"""Generate the labelled golden corpus in every supported material format.

The corpus is derived from the canonical Markdown sources in
``tests/fixtures/`` (``golden_corporate.md``, ``golden_individual.md``) so all
format variants stay in sync with the same labelled content. Generated files
are committed alongside the sources; re-run this script after editing a
source, then re-run ``scripts/evaluate_extraction.py`` and commit the updated
``docs/release/golden-report.md``.

Layout notes per format (kept deliberately simple and deterministic):

- PDF: a text-layer PDF (no OCR needed). Headings and paragraphs become native
  text lines; table rows are rendered as space-joined text lines, because a
  flat text layer cannot express the relational header structure the
  table-row extractor needs. ``evaluate_extraction.py`` therefore asserts only
  the scalar fields for the PDF fixture.
- DOCX: headings use real heading styles, key-value lines are paragraphs, and
  the financial/transaction tables are real Word tables with header rows.
- XLSX: three sheets — the material content as single-cell rows, then real
  ``科目/金额`` and ``交易日期/交易对手/金额/余额`` tables so the table-row
  extractor sees proper headers.
- CSV: single-column rows (each line one cell); like PDF, only scalar fields
  are asserted.

Usage: python scripts/generate_golden_fixtures.py
"""

import re
import sys
from dataclasses import dataclass
from pathlib import Path

import docx
import openpyxl
import pymupdf
from docx.enum.table import WD_TABLE_ALIGNMENT

FIXTURES = Path(__file__).resolve().parent.parent / "tests" / "fixtures"
OUTPUTS: tuple[tuple[str, str, list[str]], ...] = (
    ("golden_corporate", "企业贷款申请材料（黄金集）", []),
    ("golden_individual", "个人贷款申请材料（黄金集）", []),
)

# Supported structured + text-PDF formats the corpus covers.
FORMATS = (".pdf", ".docx", ".xlsx", ".csv")

# Scanned (image-only) PDF variants force the pinned PaddleOCR pipeline in
# the release golden suite (``--engine ocr``); they carry no text layer.
SCAN_FORMATS = (".scan.pdf",)

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*#*\s*$")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)*\|?\s*$")


@dataclass(frozen=True)
class Unit:
    kind: str  # "heading" | "paragraph" | "table"
    level: int = 0
    text: str = ""
    rows: tuple[tuple[str, ...], ...] = ()


def parse_source(text: str) -> list[Unit]:
    units: list[Unit] = []
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        heading = HEADING_RE.match(line)
        if heading:
            units.append(Unit("heading", len(heading.group(1)), heading.group(2).strip()))
            index += 1
            continue
        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and TABLE_SEPARATOR_RE.match(lines[index + 1])
        ):
            rows: list[tuple[str, ...]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                if TABLE_SEPARATOR_RE.match(lines[index]):
                    index += 1
                    continue
                rows.append(
                    tuple(part.strip() for part in lines[index].strip().strip("|").split("|"))
                )
                index += 1
            units.append(Unit("table", rows=tuple(rows)))
            continue
        units.append(Unit("paragraph", text=stripped))
        index += 1
    return units


def render_pdf(units: list[Unit], destination: Path) -> None:
    pdf = pymupdf.open()
    page = pdf.new_page(width=595, height=842)
    fontsize = 11
    line_height = 15
    margin = 56
    width_units = (595 - 2 * margin) / (fontsize * 0.52)
    y = margin

    def wrap(text: str) -> list[str]:
        lines: list[str] = []
        current = ""
        current_units = 0.0
        for char in text:
            units = 1.0 if ord(char) > 0x2E7F else 0.5
            if current_units + units > width_units and current:
                lines.append(current)
                current = ""
                current_units = 0.0
            current += char
            current_units += units
        if current:
            lines.append(current)
        return lines

    def emit(text: str) -> None:
        nonlocal y, page
        for line in wrap(text):
            if y > 842 - margin:
                page = pdf.new_page(width=595, height=842)
                y = margin
            page.insert_text((margin, y), line, fontname="china-s", fontsize=fontsize)
            y += line_height

    for unit in units:
        if unit.kind == "heading":
            emit(unit.text)
        elif unit.kind == "paragraph":
            emit(unit.text)
        else:
            for row in unit.rows:
                emit("  ".join(cell for cell in row if cell))
            emit("")
    pdf.save(str(destination))


def render_docx(units: list[Unit], destination: Path) -> None:
    document = docx.Document()
    for unit in units:
        if unit.kind == "heading":
            style = "Heading 1" if unit.level <= 1 else "Heading 2"
            document.add_paragraph(unit.text, style=style)
        elif unit.kind == "paragraph":
            document.add_paragraph(unit.text)
        else:
            table = document.add_table(rows=len(unit.rows), cols=len(unit.rows[0]))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for row_index, row in enumerate(unit.rows):
                for column_index, cell in enumerate(row):
                    table.cell(row_index, column_index).text = cell
            document.add_paragraph("")
    document.save(str(destination))


def split_key_value(line: str) -> list[str]:
    """Split a ``label：value`` line into a two-cell row (label, value).

    Table-row extraction reads key-value candidates from the first two columns
    of a row, while single-cell table content is only scanned for a small
    identifier whitelist — so structured single-sheet formats must present
    key-value lines as two cells to be extracted.
    """
    if "：" in line:
        label, value = line.split("：", 1)
    elif ":" in line:
        label, value = line.split(":", 1)
    else:
        return [line]
    return [label.strip(), value.strip()]


def render_xlsx(units: list[Unit], destination: Path) -> None:
    workbook = openpyxl.Workbook()
    content = workbook.active
    content.title = "材料"
    tables: list[list[Unit]] = []
    for unit in units:
        if unit.kind == "table":
            tables.append([unit])
        elif unit.kind == "heading":
            content.append([unit.text])
        else:
            content.append(split_key_value(unit.text))
    for index, table_units in enumerate(tables, start=1):
        sheet = workbook.create_sheet(title=f"表格{index}")
        for unit in table_units:
            for row in unit.rows:
                sheet.append(list(row))
    workbook.save(str(destination))


def render_csv(units: list[Unit], destination: Path) -> None:
    rows: list[list[str]] = []
    for unit in units:
        if unit.kind == "table":
            for row in unit.rows:
                rows.append(list(row))
        else:
            rows.append(split_key_value(unit.text))
    destination.write_text("\n".join(",".join(row) for row in rows) + "\n", encoding="utf-8")


def render_scan_pdf(units: list[Unit], destination: Path) -> None:
    """Render the content to an image-only PDF (no text layer).

    Mirrors the worker's OCR path: each page is rasterized at 2x and embedded
    as an image, so ``parse_material`` runs the pinned PaddleOCR models
    (``run_ocr=True``) instead of reading a text layer.
    """
    text_pdf = pymupdf.open()
    page = text_pdf.new_page(width=595, height=842)
    fontsize = 11
    line_height = 15
    margin = 56
    width_units = (595 - 2 * margin) / (fontsize * 0.52)
    y = margin

    def wrap(text: str) -> list[str]:
        lines: list[str] = []
        current = ""
        current_units = 0.0
        for char in text:
            char_units = 1.0 if ord(char) > 0x2E7F else 0.5
            if current_units + char_units > width_units and current:
                lines.append(current)
                current = ""
                current_units = 0.0
            current += char
            current_units += char_units
        if current:
            lines.append(current)
        return lines

    def emit(text: str) -> None:
        nonlocal y, page
        for line in wrap(text):
            if y > 842 - margin:
                page = text_pdf.new_page(width=595, height=842)
                y = margin
            page.insert_text((margin, y), line, fontname="china-s", fontsize=fontsize)
            y += line_height

    for unit in units:
        if unit.kind == "heading":
            emit(unit.text)
        elif unit.kind == "paragraph":
            emit(unit.text)
        else:
            for row in unit.rows:
                emit("  ".join(cell for cell in row if cell))
            emit("")
    scan = pymupdf.open()
    for page in text_pdf:
        pixmap = page.get_pixmap(matrix=pymupdf.Matrix(2, 2), alpha=False)
        scan_page = scan.new_page(width=595, height=842)
        image = pixmap.tobytes("png")
        scan_page.insert_image(pymupdf.Rect(0, 0, 595, 842), stream=image)
    scan.save(str(destination), deflate=True, garbage=4)
    text_pdf.close()
    scan.close()


def main() -> None:
    generated: list[str] = []
    for name, title, _ in OUTPUTS:
        source = (FIXTURES / f"{name}.md").read_text(encoding="utf-8")
        units = parse_source(source)
        for extension in FORMATS:
            destination = FIXTURES / f"{name}{extension}"
            if extension == ".pdf":
                render_pdf(units, destination)
            elif extension == ".docx":
                render_docx(units, destination)
            elif extension == ".xlsx":
                render_xlsx(units, destination)
            elif extension == ".csv":
                render_csv(units, destination)
            generated.append(str(destination))
        for extension in SCAN_FORMATS:
            destination = FIXTURES / f"{name}{extension}"
            render_scan_pdf(units, destination)
            generated.append(str(destination))
    print(f"generated {len(generated)} fixtures:")
    for path in generated:
        print(f"  {path}")


if __name__ == "__main__":
    sys.exit(main())
